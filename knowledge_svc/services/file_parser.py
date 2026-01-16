from typing import Optional
import io
from pathlib import Path

# PDF parsing
import PyPDF2
import pdfplumber

# DOCX parsing
from docx import Document

# Markdown parsing
import markdown
from html.parser import HTMLParser


class MLStripper(HTMLParser):
    """Helper to strip HTML tags from markdown conversion."""
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []
    
    def handle_data(self, d):
        self.text.append(d)
    
    def get_data(self):
        return ''.join(self.text)


def strip_html_tags(html: str) -> str:
    """Strip HTML tags from string."""
    s = MLStripper()
    s.feed(html)
    return s.get_data()


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2 and pdfplumber."""
    text_parts = []
    
    try:
        # Try pdfplumber first (better for tables and layout)
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        print(f"pdfplumber failed, trying PyPDF2: {e}")
        try:
            # Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        except Exception as e2:
            print(f"PyPDF2 also failed: {e2}")
            raise ValueError(f"Failed to parse PDF: {e2}")
    
    return "\n\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")


def parse_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file."""
    try:
        # Try UTF-8 first
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        try:
            # Fallback to latin-1
            return file_bytes.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {e}")


def parse_markdown(file_bytes: bytes) -> str:
    """Convert Markdown to plain text."""
    try:
        md_text = file_bytes.decode('utf-8')
        # Convert to HTML then strip tags
        html = markdown.markdown(md_text)
        plain_text = strip_html_tags(html)
        return plain_text
    except Exception as e:
        raise ValueError(f"Failed to parse Markdown: {e}")


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension."""
    return Path(filename).suffix.lower()


def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Parse file based on extension.
    
    Args:
        filename: Name of the file
        file_bytes: Raw file bytes
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is unsupported or parsing fails
    """
    ext = get_file_extension(filename)
    
    parsers = {
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.doc': parse_docx,  # Try docx parser for .doc too
        '.txt': parse_txt,
        '.md': parse_markdown,
        '.markdown': parse_markdown,
    }
    
    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(parsers.keys())}")
    
    try:
        text = parser(file_bytes)
        if not text or not text.strip():
            raise ValueError(f"No text content extracted from {filename}")
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error parsing {filename}: {e}")
